import os
from typing import List, Dict
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from documents.document import Document


def load_pdf(file_path: str) -> str:
    """Load content from a PDF file."""
    try:
        with open(file_path, "rb") as file:
            pdf_reader = PdfReader(file)
            content = ""
            for page in pdf_reader.pages:
                content += page.extract_text() + "\n"
        return content
    except ImportError:
        raise ValueError(
            "PyPDF2 library is required for PDF files. Install with: pip install PyPDF2"
        )


def load_docx(path: str) -> str:
    """Load a .docx file and return its full text."""
    doc = DocxDocument(path)
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])


def load_txt(file_path: str) -> str:
    """Load content from a text file."""
    try:
        with open(file_path, "r", encoding="utf-8") as file:
            content = file.read()
        return content
    except UnicodeDecodeError:
        try:
            with open(file_path, "r", encoding="latin-1") as file:
                content = file.read()
            return content
        except Exception as e:
            raise ValueError(f"Error reading text file: {e}")


def load_document(file_path) -> List[Dict[str, str]]:
    """
    Load the document based on file type.

    Args:
        file_path (str): The path to the file to load.

    Returns:
        List[Dict[str, str]]: A list of dictionaries with 'page_content' key containing the file content.

    Raises:
        ValueError: If the file type is not supported.
    """

    text_extensions = [".txt", ".md", ".py", ".yaml", ".yml", ".json"]

    ext = os.path.splitext(file_path)[-1].lower()
    content = ""

    if ext == ".pdf":
        content = load_pdf(file_path)
    elif ext == ".docx":
        content = load_docx(file_path)
    elif ext in text_extensions:
        content = load_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    return Document(content, metadata={"file_path": file_path})
